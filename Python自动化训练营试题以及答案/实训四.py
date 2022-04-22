"""
  1、提取出公司合同word中，甲方和乙方的银行账号，用字典进行存储，
  例如{‘甲方’：123,‘乙方’：456}
"""
import os
# docx 处理模块
from collections import defaultdict
import docx
import re

'''do_path的like版本'''


def do_path_like(path, type):
    file_list = []
    for root, dirs, files in os.walk(path):
        for name in files:
            # 通过增加if 模式，增加like匹配
            if type in name:
                file_list.append(os.path.join(root, name))
    # print(file_list)
    return file_list


# def word_to_text(path, file_type):
#     word_list = do_path_like(path, file_type)
#     for i in word_list:
#         doc = docx.Document(i)
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     print(cell.text)


def word_to_text(path, file_type):
    word_list = do_path_like(path, file_type)

    # 多个表
    docx_tbs = [tb for i in word_list for tb in docx.Document(i).tables]

    # 多个单元行
    rows = (row for tb in docx_tbs for row in tb.rows)
    # 多个单元格
    cells = (cell for row in rows for cell in row.cells)

    cp_name = r"(.)方.*：(.*)"
    zh_re = r"账.*：(.*)"

    target_dict = defaultdict(list)  # 存储的目标字典

    for i in cells:
        cp_txt = re.search(cp_name, i.text)
        zh_txt = re.search(zh_re, i.text)
        if cp_txt and zh_txt:
            # 公司名字
            cp_name_res = cp_txt.group(1) + '-' + cp_txt.group(2).strip()
            # 银行账号
            zh_res = zh_txt.group(1)

            target_dict[cp_name_res].append(zh_res)

    return target_dict


path_file = os.path.abspath(os.path.dirname(__file__))

# target = word_to_text(path_file, '合同')
# print(target)

"""
  2、批量发送邮件，通过yagmail ，给自己的3位QQ好友，发送过去3封不同的邮件内容。
"""
import yagmail
import requests
import time
from PIL import Image


def do_email():
    yag = yagmail.SMTP(user="ityige@126.com", password="ityige666", host='smtp.126.com')
    # yag = yagmail.SMTP(user="1071235258@qq.com", password="npwavzdmabffbedh", host='smtp.qq.com')
    send_dujt = 'https://api.oick.cn/dutang/api.php'
    send_bg = 'https://api.oick.cn/random/api.php'

    qq_message = [1071235258, 2369297366, 429614073]

    for index, qq in enumerate(qq_message):
        mail = str(qq) + '@qq.com'
        res_text = requests.get(send_dujt).text

        res_jpg = requests.get(send_bg)

        # 确定图片的存储地址
        img_path = os.getcwd() + '\\images\\' + str(index) + '.jpg'

        # 图片进行存储
        with open(img_path, 'wb') as f:
            f.write(res_jpg.content)

        image = Image.open(img_path)
        resized_image = image.resize((800, 500), Image.ANTIALIAS)
        resized_image.save(img_path)

        # 发送的内容
        contents = [res_text, yagmail.inline(img_path)]

        subject = time.strftime("%Y-%m-%d") + '-毒鸡汤'

        yag.send(to=mail, subject=subject, contents=contents)

        print(mail, ':发送完毕！')

do_email()


####需求6：按照北京的区域，分隔成若干个Excel文件###########################################################
import pandas as pd
houses = pd.read_excel("D:\\rd\PyJob\Python自动化办公训练营\Python自动化训练营试题以及答案\house.xlsx")

listType = houses['朝向'].unique()
print(listType) #< class 'numpy.ndarray'>

list_res = list(listType)

print('带有nan的列表',list_res) # ['朝南', '朝西', '朝东', '朝北', nan]

list_res.remove(np.nan)

while np.nan in list_res :
    list_res.remove(np.nan)
print('处理好的列表',list_res)
